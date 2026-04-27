"""
Génère le Cahier des Charges BloodLink v2.0 en format .docx
avec les diagrammes Mermaid convertis en images PNG.
"""

import re
import io
import requests
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

# ── Chemins ──────────────────────────────────────────────
BASE = Path(r"c:\Users\achar\Documents\Boold_link\docs")
MD_FILE = BASE / "CAHIER_DES_CHARGES_V2.md"
OUT_DOCX = BASE / "docx" / "CAHIER_DES_CHARGES_BloodLink_v2.docx"
MERMAID_DIR = BASE / "mermaid_imgs"

MERMAID_DIR.mkdir(parents=True, exist_ok=True)
OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

# ── Lire le markdown ────────────────────────────────────
md_text = MD_FILE.read_text(encoding="utf-8")

# ── Extraire les blocs Mermaid ───────────────────────────
mermaid_blocks = re.findall(r'```mermaid\n(.*?)```', md_text, re.DOTALL)
mermaid_names = [
    "use_case_diagram",
    "architecture_diagram",
    "uml_class_diagram",
    "er_diagram",
    "sequence_diagram",
]

# ── Générer les images Mermaid via mermaid.ink ───────────
print("Génération des diagrammes Mermaid...")
mermaid_images = {}

for i, code in enumerate(mermaid_blocks):
    name = mermaid_names[i] if i < len(mermaid_names) else f"diagram_{i}"
    img_path = MERMAID_DIR / f"{name}.png"

    if img_path.exists():
        print(f"  ✓ {name} (cache)")
        mermaid_images[i] = img_path
        continue

    # Utiliser l'API mermaid.ink
    try:
        import base64
        encoded = base64.urlsafe_b64encode(code.encode('utf-8')).decode('ascii')
        url = f"https://mermaid.ink/img/{encoded}"
        resp = requests.get(url, timeout=30)
        if resp.status_code == 200:
            img_path.write_bytes(resp.content)
            print(f"  ✓ {name}")
            mermaid_images[i] = img_path
        else:
            print(f"  ✗ {name} - HTTP {resp.status_code}")
    except Exception as e:
        print(f"  ✗ {name} - {e}")

# ── Créer le document Word ──────────────────────────────
print("Création du document Word...")
doc = Document()

# ── Styles ───────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Titre principal
title = doc.add_heading('Cahier des Charges — Projet BloodLink (MVP) v2.0', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0xB2, 0x22, 0x22)  # Rouge sang

doc.add_paragraph(
    'Application de mise en relation entre donneurs de sang et centres de dons.',
    style='Subtitle'
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── Parser le markdown et créer le docx ──────────────────
lines = md_text.split('\n')
i = 0
mermaid_idx = 0
in_table = False
table_rows = []

def add_table_from_rows(doc, rows):
    """Ajoute un tableau Word à partir d'une liste de listes de strings."""
    if not rows:
        return
    # Nettoyer les cellules
    cleaned = []
    for row in rows:
        cleaned_row = [cell.strip() for cell in row]
        cleaned.append(cleaned_row)

    ncols = max(len(r) for r in cleaned)
    # Compléter les lignes plus courtes
    for r in cleaned:
        while len(r) < ncols:
            r.append('')

    table = doc.add_table(rows=len(cleaned), cols=ncols, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row in enumerate(cleaned):
        for ci, cell in enumerate(row):
            cell_obj = table.cell(ri, ci)
            # Gras pour la première ligne (en-tête)
            if ri == 0:
                run = cell_obj.paragraphs[0].add_run(cell)
                run.bold = True
                run.font.size = Pt(10)
            else:
                run = cell_obj.paragraphs[0].add_run(cell)
                run.font.size = Pt(10)
    doc.add_paragraph()  # espace après tableau

def parse_table_row(line):
    """Parse une ligne de tableau markdown en liste de strings."""
    cells = line.strip().split('|')
    # Retirer le premier et dernier élément vide (dû aux | de début/fin)
    cells = [c.strip() for c in cells if c.strip() != '' or True]
    # Meilleur parsing
    if line.strip().startswith('|'):
        cells = line.strip().split('|')
        cells = [c.strip() for c in cells[1:-1]]  # enlever les bords
    return cells

while i < len(lines):
    line = lines[i]

    # ── Sauter les lignes de séparation ──
    if line.strip() == '---':
        i += 1
        continue

    # ── Tableau markdown ──
    if '|' in line and line.strip().startswith('|'):
        # Collecter toutes les lignes du tableau
        table_rows = []
        while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
            row_line = lines[i].strip()
            # Ignorer les lignes de séparation du tableau (|---|---|)
            if re.match(r'^\|[\s\-:|]+\|$', row_line):
                i += 1
                continue
            cells = parse_table_row(row_line)
            table_rows.append(cells)
            i += 1
        add_table_from_rows(doc, table_rows)
        continue

    # ── Bloc Mermaid ──
    if line.strip().startswith('```mermaid'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        i += 1  # sauter le ``` fermant

        # Ajouter l'image si disponible
        if mermaid_idx in mermaid_images:
            img_path = mermaid_images[mermaid_idx]
            name = mermaid_names[mermaid_idx] if mermaid_idx < len(mermaid_names) else f"diagram_{mermaid_idx}"
            doc.add_paragraph(f"Diagramme : {name}", style='Caption')
            try:
                doc.add_picture(str(img_path), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Image non disponible : {e}]")
        else:
            # Ajouter le code en bloc citation
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text, style='No Spacing')
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        doc.add_paragraph()
        mermaid_idx += 1
        continue

    # ── Autres blocs de code ──
    if line.strip().startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        i += 1
        code_text = '\n'.join(code_lines)
        p = doc.add_paragraph(code_text, style='No Spacing')
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        doc.add_paragraph()
        continue

    # ── Titres ──
    if line.startswith('# ') and not line.startswith('## '):
        # Titre principal déjà ajouté
        i += 1
        continue
    elif line.startswith('## '):
        heading_text = line.replace('## ', '').strip()
        doc.add_heading(heading_text, level=1)
    elif line.startswith('### '):
        heading_text = line.replace('### ', '').strip()
        doc.add_heading(heading_text, level=2)
    elif line.startswith('#### '):
        heading_text = line.replace('#### ', '').strip()
        doc.add_heading(heading_text, level=3)

    # ── Citations (> ...) ──
    elif line.strip().startswith('>'):
        quote_text = line.strip().lstrip('> ').strip()
        p = doc.add_paragraph(quote_text, style='Quote')
        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Listes à puces ──
    elif line.strip().startswith('- [ ]') or line.strip().startswith('- [x]'):
        text = re.sub(r'^- \[[ x]\] ', '', line.strip())
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        # Enlever le markdown bold
        run.bold = '**' in text
        text_clean = text.replace('**', '')
        p.clear()
        # Gérer le bold dans le texte
        parts = text_clean.split('**') if '**' in text else [text_clean]
        # Simplification: juste ajouter le texte nettoyé
        p.add_run(text_clean.replace('**', ''))
    elif line.strip().startswith('- '):
        text = line.strip().lstrip('- ').strip()
        text = text.replace('**', '')
        text = text.replace('`', '')
        doc.add_paragraph(text, style='List Bullet')

    # ── Lignes normales ──
    elif line.strip():
        text = line.strip()
        # Gérer le bold et inline code
        p = doc.add_paragraph()
        # Parser les segments bold et code
        segments = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
        for seg in segments:
            if seg.startswith('**') and seg.endswith('**'):
                run = p.add_run(seg[2:-2])
                run.bold = True
            elif seg.startswith('`') and seg.endswith('`'):
                run = p.add_run(seg[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            else:
                p.add_run(seg)

    i += 1

# ── Sauvegarder ──────────────────────────────────────────
doc.save(str(OUT_DOCX))
print(f"\n✅ Document Word généré : {OUT_DOCX}")
print(f"   Diagrammes Mermaid : {len(mermaid_images)}/{len(mermaid_blocks)} générés")
